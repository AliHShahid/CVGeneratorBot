import os
import re
from typing import Dict, List, Optional
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


def create_output_directory(directory: str = "output") -> None:
    """Create output directory if it doesn't exist."""
    os.makedirs(directory, exist_ok=True)


def validate_email(email: str) -> bool:
    """Validate email format."""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None


def validate_phone(phone: str) -> bool:
    """Validate phone number format (German format)."""
    # Clean phone number
    clean_phone = re.sub(r'[^\d+]', '', phone)
    # Check if it's a valid German phone number format
    patterns = [
        r'^\+49\d{10,11}$',  # +49 followed by 10-11 digits
        r'^0\d{9,11}$',      # 0 followed by 9-11 digits
    ]
    return any(re.match(pattern, clean_phone) for pattern in patterns)


def clean_text(text: str) -> str:
    """Clean and normalize text."""
    if not text:
        return ""
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s@.+-]', '', text)
    return text.strip()


def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information from text using regex patterns."""
    contact_info = {
        'email': '',
        'phone': '',
        'name': ''
    }
    
    # Email extraction
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        contact_info['email'] = emails[0]
    
    # Phone extraction (German format)
    phone_patterns = [
        r'\+49[-\s]?\d{3}[-\s]?\d{3}[-\s]?\d{4,5}',
        r'0\d{3}[-\s]?\d{3}[-\s]?\d{4,5}',
        r'\+49[-\s]?\d{10,11}',
        r'0\d{9,11}'
    ]
    
    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        if phones:
            contact_info['phone'] = phones[0]
            break
    
    return contact_info


def extract_skills_keywords(text: str) -> List[str]:
    """Extract potential skills and keywords from text."""
    # Common German skill keywords
    skill_patterns = [
        r'\b(?:Python|Java|JavaScript|C\+\+|SQL|HTML|CSS)\b',
        r'\b(?:SAP|Workday|MS Office|Excel|PowerPoint)\b',
        r'\b(?:Projektmanagement|Führung|Kommunikation|Teamwork)\b',
        r'\b(?:Deutsch|Englisch|Französisch|Spanisch)\b',
    ]
    
    skills = []
    for pattern in skill_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        skills.extend(matches)
    
    return list(set(skills))  # Remove duplicates


def format_profile_text(profile_data: Dict) -> str:
    """Format profile data into readable text according to German template."""
    lines = []
    
    # Header section
    lines.append("**Titel des Job Postings**")
    lines.append("")
    lines.append("| | **angefragt** | **falls abweichend** |")
    lines.append("|---|---|---|")
    lines.append(f"| **Einkaufskurzprofil (EKP)** | | {profile_data.get('einkaufskurzprofil', '')} |")
    lines.append(f"| **Stundenverrechnungssatz (SVS)** | **{profile_data.get('stundenverrechnungssatz', '€')}** | **€** |")
    lines.append(f"| **Möglicher Starttermin** | {profile_data.get('starttermin', '')} | |")
    lines.append("")
    
    # Professional Experience
    lines.append("**Berufserfahrung:**")
    lines.append("")
    for exp in profile_data.get('berufserfahrung', []):
        lines.append(f"**{exp.get('period', '')}** | {exp.get('company', '')}")
        lines.append(f"{exp.get('description', '')}")
        lines.append("")
    
    # Education
    lines.append("**Ausbildung:**")
    lines.append("")
    for edu in profile_data.get('ausbildung', []):
        lines.append(f"**{edu.get('period', '')}** | {edu.get('institution', '')}")
        lines.append(f"{edu.get('description', '')}")
        lines.append("")
    
    # Skills
    lines.append("**Kompetenzen:**")
    lines.append("")
    lines.append("EDV-Kenntnisse:")
    lines.append("")
    for skill in profile_data.get('edv_kenntnisse', []):
        lines.append(f"| {skill}: | Advanced |")
    lines.append("")
    
    lines.append("Sonstige Techniken:")
    lines.append("")
    for skill in profile_data.get('sonstige_techniken', []):
        lines.append(f"| {skill}: | Advanced |")
    lines.append("")
    
    lines.append("Sprachkenntnisse:")
    lines.append("")
    for lang in profile_data.get('sprachkenntnisse', []):
        lines.append(f"| {lang}: | Advanced |")
    lines.append("")
    
    # Additional remarks
    lines.append("**Zusätzliche Bemerkungen**")
    lines.append("")
    lines.append(profile_data.get('zusaetzliche_bemerkungen', ''))
    
    return '\n'.join(lines)

def create_german_cv_docx(profile_data: Dict, filename: str) -> str:
    """Create a DOCX file matching the exact German CV template format."""
    create_output_directory()
    filepath = os.path.join("output", filename)
    
    doc = Document()
    
    # Add the header table with 4 rows and 3 columns
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # First row - Title
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Titel des Job Postings"
    hdr_cells[1].text = ""
    hdr_cells[2].text = ""
    
    # Second row - Subheaders
    sub_cells = table.rows[1].cells
    sub_cells[0].text = ""
    sub_cells[1].text = "angefragt"
    sub_cells[2].text = "falls abweichend"
    
    # Third row - EKP
    ekp_cells = table.rows[2].cells
    ekp_cells[0].text = "Einkaufskurzprofil (EKP)"
    ekp_cells[1].text = ""
    ekp_cells[2].text = profile_data.get('einkaufskurzprofil', 'X|YYY|XXX|Z')
    
    # Fourth row - SVS
    svs_cells = table.rows[3].cells
    svs_cells[0].text = "Stundenverrechnungssatz (SVS)"
    svs_cells[1].text = profile_data.get('stundenverrechnungssatz', '€')
    svs_cells[2].text = "€"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add possible start date
    start_table = doc.add_table(rows=1, cols=3)
    start_table.style = 'Table Grid'
    start_cells = start_table.rows[0].cells
    start_cells[0].text = "Möglicher Starttermin (verfügbar ab / Kündigungsfrist)"
    start_cells[1].text = profile_data.get('starttermin', '01.04.2025')
    start_cells[2].text = ""
    
    # Add space
    doc.add_paragraph()
    
    # Add personal data section
    doc.add_heading("Persönliche Daten des Zeitarbeitnehmers", level=2)
    
    # Create personal data table
    personal_table = doc.add_table(rows=7, cols=2)
    personal_table.style = 'Table Grid'
    
    # Fill personal data
    personal_cells = personal_table.rows[0].cells
    personal_cells[0].text = "Anrede"
    personal_cells[1].text = ""
    
    personal_cells = personal_table.rows[1].cells
    personal_cells[0].text = "Vorname"
    personal_cells[1].text = profile_data.get('vorname', '')
    
    personal_cells = personal_table.rows[2].cells
    personal_cells[0].text = "Nachname"
    personal_cells[1].text = profile_data.get('nachname', '')
    
    personal_cells = personal_table.rows[3].cells
    personal_cells[0].text = "Geburtsdatum"
    personal_cells[1].text = profile_data.get('geburtsdatum', '')
    
    personal_cells = personal_table.rows[4].cells
    personal_cells[0].text = "Geburtsort"
    personal_cells[1].text = profile_data.get('geburtsort', '')
    
    personal_cells = personal_table.rows[5].cells
    personal_cells[0].text = "Staatsangehörigkeit"
    personal_cells[1].text = profile_data.get('staatsangehoerigkeit', '')
    
    # Add space
    doc.add_paragraph()
    
    # Add work experience section
    doc.add_heading("Berufserfahrung:", level=2)
    
    for exp in profile_data.get('berufserfahrung', []):
        # Create experience table with 1 row and 2 columns
        exp_table = doc.add_table(rows=1, cols=2)
        exp_table.style = 'Table Grid'
        
        # First column - period
        period_cell = exp_table.cell(0, 0)
        period_cell.text = exp.get('period', '')
        
        # Second column - company and description
        desc_cell = exp_table.cell(0, 1)
        desc_cell.text = f"{exp.get('company', '')}\n\n{exp.get('description', '')}"
        
        # Add space after each experience
        doc.add_paragraph()
    
    # Add education section
    doc.add_heading("Ausbildung:", level=2)
    
    for edu in profile_data.get('ausbildung', []):
        # Create education table with 1 row and 2 columns
        edu_table = doc.add_table(rows=1, cols=2)
        edu_table.style = 'Table Grid'
        
        # First column - period
        period_cell = edu_table.cell(0, 0)
        period_cell.text = edu.get('period', '')
        
        # Second column - institution and description
        desc_cell = edu_table.cell(0, 1)
        desc_cell.text = f"{edu.get('institution', '')}\n\n{edu.get('description', '')}"
        
        # Add space after each education entry
        doc.add_paragraph()
    
    # Add skills section
    doc.add_heading("Kompetenzen:", level=2)
    
    # Add IT skills
    doc.add_paragraph("EDV-Kenntnisse:")
    for skill in profile_data.get('edv_kenntnisse', []):
        skill_table = doc.add_table(rows=1, cols=2)
        skill_table.style = 'Table Grid'
        skill_cell = skill_table.cell(0, 0)
        skill_cell.text = f"{skill}:"
        level_cell = skill_table.cell(0, 1)
        level_cell.text = "Advanced"
    
    # Add space
    doc.add_paragraph()
    
    # Add other technical skills
    doc.add_paragraph("Sonstige Techniken:")
    for skill in profile_data.get('sonstige_techniken', []):
        skill_table = doc.add_table(rows=1, cols=2)
        skill_table.style = 'Table Grid'
        skill_cell = skill_table.cell(0, 0)
        skill_cell.text = f"{skill}:"
        level_cell = skill_table.cell(0, 1)
        level_cell.text = "Advanced"
    
    # Add space
    doc.add_paragraph()
    
    # Add languages
    doc.add_paragraph("Sprachkenntnisse:")
    for lang in profile_data.get('sprachkenntnisse', []):
        lang_table = doc.add_table(rows=1, cols=2)
        lang_table.style = 'Table Grid'
        lang_cell = lang_table.cell(0, 0)
        lang_cell.text = f"{lang}:"
        level_cell = lang_table.cell(0, 1)
        level_cell.text = "Advanced" if lang != "Hindi" else "Expert"
    
    # Add space
    doc.add_paragraph()
    
    # Add additional remarks
    doc.add_heading("Zusätzliche Bemerkungen", level=2)
    remarks_table = doc.add_table(rows=1, cols=1)
    remarks_table.style = 'Table Grid'
    remarks_cell = remarks_table.cell(0, 0)
    remarks_cell.text = profile_data.get('zusaetzliche_bemerkungen', '')
    
    doc.save(filepath)
    return filepath

def save_text_to_file(text: str, filename: str, directory: str = "output") -> str:
    """Save text to a file."""
    create_output_directory(directory)
    filepath = os.path.join(directory, filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(text)
    
    return filepath


def read_pdf_content(pdf_path: str) -> str:
    """Read content from PDF file."""
    try:
        with fitz.open(pdf_path) as doc:
            text = ""
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""


def read_docx_content(docx_path: str) -> str:
    """Read content from DOCX file."""
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""


def generate_unique_filename(base_name: str, extension: str) -> str:
    """Generate a unique filename with timestamp."""
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base_name}_{timestamp}.{extension}"


def log_processing_info(text_length: int, processing_time: float) -> None:
    """Log processing information."""
    print(f"Processed {text_length} characters in {processing_time:.2f} seconds")